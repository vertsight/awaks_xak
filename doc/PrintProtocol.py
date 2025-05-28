from typing import Any, Dict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import json
import os

from doc.LoadData import load_protocol_data

def create_protocol(conference_id: int, db_config: Dict[str, str]) -> Dict[str, Any]:
    data = load_protocol_data(conference_id, db_config)
    json_file_path = os.path.join("Data", "obtainedReportData.json")

    if not os.path.exists(json_file_path):
        raise FileNotFoundError(f"JSON file not found at {json_file_path}")

    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        raise Exception(f"Error reading JSON from {json_file_path}: {e}")

    doc = Document()

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Cm(1.25)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ПРОТОКОЛ")
    run.bold = True
    run.font.size = Pt(18)

    meeting_date = datetime.strptime(data['meeting_date'], '%Y-%m-%d').strftime('%d-%m-%Y')
    date_title = doc.add_paragraph()
    date_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_title.add_run(f"Документ от {meeting_date}")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    attendees_title = doc.add_paragraph()
    run = attendees_title.add_run("На совещание присутствуют следующие лица:")
    run.bold = True
    for person in data['attendees']:
        p = doc.add_paragraph(f"- {person['name']} ({person['position']});")
    doc.add_paragraph()

    topics_title = doc.add_paragraph()
    run = topics_title.add_run("Темы для дискуссии:")
    run.bold = True
    for i, topic in enumerate(data['topics'], 1):
        p = doc.add_paragraph(f"{i}. {topic['title']} - {topic['description']};")
    doc.add_paragraph()

    # decisions_title = doc.add_paragraph()
    # run = decisions_title.add_run("Решения вынесенные в ходе обсуждения:")
    # run.bold = True
    # for decision in data['decisions']:
    #     topic_title = data['topics'][decision['topic_index']]['title']
    #     p = doc.add_paragraph(f"- По теме {decision['topic_index'] + 1} \"{topic_title}\" - {decision['decision']};")
    # doc.add_paragraph()

    # responsibles_title = doc.add_paragraph()
    # run = responsibles_title.add_run("Ответственные за реализацию или управления поставленными решениями:")
    # run.bold = True
    # for responsible in data['responsibles']:
    #     topic_title = data['topics'][responsible['topic_index']]['title']
    #     p = doc.add_paragraph(
    #         f"- Выполнением решения темы {responsible['topic_index'] + 1} \"{topic_title}\", "
    #         f"назначен: {responsible['name']} ({responsible['position']}), "
    #         f"в качестве выполняемого обязанности {responsible['responsibilities']};"
    #     )
    # doc.add_paragraph()

    doc.add_paragraph()
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    for cell in table.columns[0].cells:
        cell.width = Cm(8)
    for cell in table.columns[1].cells:
        cell.width = Cm(8)
    
    left_cell = table.cell(0, 0)
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_paragraph.add_run("Должность:")
    left_run.bold = True
    
    right_cell = table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right_run = right_paragraph.add_run("Ф.И.О.:")
    right_run.bold = True
    
    doc.add_paragraph()
    sign_paragraph = doc.add_paragraph()
    sign_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_run = sign_paragraph.add_run("_________________/_________________")
    sign_run.font.size = Pt(12)
    
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_paragraph.add_run("Дата: _______________")
    date_run.font.size = Pt(12)

    print(data['output_filename'])
    doc.save(data['output_filename'])
    filename = data['output_filename']
    full_path = os.path.abspath(filename)
    return full_path